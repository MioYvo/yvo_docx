# __author__ = "Mio"
# __email__: "liurusi.101@gmail.com"
# created: 5/24/21 11:01 PM
from enum import Enum
from pathlib import Path
from typing import List, Optional

from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Cm, Pt
from docx.table import _Cell, _Row, Table
from docx.text.paragraph import Paragraph

from docx import Document
from docx.document import Document as T_Document
from docx.text.run import Run
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.oxml.ns import qn


class ProcessType(Enum):
    one_para = 'one_para'
    name_para = 'name_para'
    multi_para = 'multi_para'
    one_multi_para = 'one_multi_para'
    table = 'table'
    runs = 'runs'


class P:
    def __init__(self,
                 name: str,
                 num: int = 0,
                 old_startswith: str = None,
                 waiting: bool = False,
                 parent_p: "P" = None,
                 process_type: ProcessType = ProcessType.runs,
                 table_first_cell_text: str = '',
                 need_wait: bool = True,
                 multi_para_multi_runs: bool = True, one_para_contain_p_name: bool = False):
        self.name = name
        self.num = num
        self.old_startswith: str = old_startswith or name
        self.table_first_cell_text = table_first_cell_text
        self.one_para_contain_p_name = one_para_contain_p_name
        self.is_done = False
        self.waiting = waiting
        self.sub_p: List["P"] = list()
        self.parent_p = parent_p
        if self.parent_p:
            self.parent_p.add_sub_p(self)
        self.need_wait = need_wait
        self.multi_para_multi_runs = multi_para_multi_runs

        self.process_type: ProcessType = process_type

    def add_sub_p(self, p: "P"):
        self.sub_p.append(p)

    def check_sub_p_done(self):
        if set(_p.is_done for _p in self.sub_p) == {True}:
            self.done()

    def done(self):
        self.waiting = False
        self.is_done = True
        if self.parent_p:
            self.parent_p.check_sub_p_done()


class LoopControl(Enum):
    continue_ = "Continue"
    break_ = "Break"


def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, T_Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    elif isinstance(parent, _Row):
        parent_elm = parent._tr
    else:
        raise ValueError("something's not right")
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


class YvoParser:
    def __init__(self,
                 old_doc=Path()/'重铬酸钾__Potassium dichromate__7778-50-9.docx',
                 new_doc_path: Path = Path('sds_v1/')):
        self.old_doc: T_Document = Document(old_doc)
        self.old_doc_name = old_doc.name
        self.new_doc: T_Document = Document()
        sections = self.new_doc.sections
        for section in sections:
            section.top_margin = Cm(0.5)
            section.bottom_margin = Cm(0.5)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)

        self.new_doc.styles['Normal'].font.name = u'宋体'
        self.new_doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        self.new_doc_path = new_doc_path
        self.new_doc.add_heading('化学品安全技术说明书', 0)

        self.multi_para_multi_runs_p: Optional[P] = None

        self.p1 = P('化学品及企业标识', num=1, old_startswith='第1部分 化学品及企业标识', process_type=ProcessType.name_para,
                    need_wait=False, one_para_contain_p_name=True)   # 化学品及企业标识

        self.p1_1 = P('化学品中文名', parent_p=self.p1)
        self.p1_2 = P('化学品英文名', parent_p=self.p1)

        self.p2 = P('危险性概述', num=2, old_startswith='第2部分 危险性概述', process_type=ProcessType.multi_para)   # '危险性概述'
        self.p21 = P('紧急情况概述', parent_p=self.p2)  # '紧急情况概述'
        self.p22 = P('废弃处置', parent_p=self.p2)     # '废弃处置'

        self.p3 = P('成分/组成信息', num=3, old_startswith='第3部分 成分/组成信息', process_type=ProcessType.table,
                    table_first_cell_text='组分')

        self.p4 = P('急救措施', num=4, old_startswith='第4部分 急救措施', process_type=ProcessType.one_para)

        self.p5 = P('消防措施', num=5, old_startswith='第5部分 消防措施', process_type=ProcessType.multi_para)
        self.p51 = P('灭火剂', parent_p=self.p5)
        self.p51 = P('特别危险性', parent_p=self.p5)
        self.p51 = P('灭火注意事项及防护措施', parent_p=self.p5)

        self.p6 = P('泄露应急处理', num=6, old_startswith='第6部分 泄露应急处理', process_type=ProcessType.multi_para)
        self.p61 = P('作业人员防护措施、防护装备和应急处置程序', parent_p=self.p6)
        self.p62 = P('泄漏化学品的收容、清除方法及所使用的处置材料', parent_p=self.p6)

        self.p7 = P('操作处置与储存', num=7, old_startswith='第7部分 操作处置与储存', process_type=ProcessType.multi_para)
        self.p71 = P('操作注意事项', parent_p=self.p7)
        self.p71 = P('储存注意事项', parent_p=self.p7)

        self.p8 = P('接触控制/个体防护', num=8, old_startswith='第8部分 接触控制/个体防护', process_type=ProcessType.multi_para)
        self.p81 = P('工程控制', parent_p=self.p8)
        self.p82 = P('个体防护装备', parent_p=self.p8)

        self.p9 = P('理化特性', num=9, old_startswith='第9部分 理化特性', process_type=ProcessType.table,
                    table_first_cell_text='外观与性状')

        self.p10 = P('稳定性和反应性', num=10, old_startswith='第10部分 稳定性和反应性', process_type=ProcessType.one_para,
                     need_wait=False, one_para_contain_p_name=True)

        self.p11 = P('废弃处置', num=11, old_startswith='第13部分 废弃处置', process_type=ProcessType.multi_para)
        self.p11_1 = P('废弃化学品', parent_p=self.p11)
        self.p11_2 = P('污染包装物', parent_p=self.p11)
        self.p11_3 = P('废弃注意事项', parent_p=self.p11)

        self.p12 = P('运输信息', num=12, old_startswith='第14部分 运输信息', process_type=ProcessType.one_para,
                     need_wait=False, one_para_contain_p_name=True)

        # self.p12_1 = P('运输注意事项', parent_p=self.p12)

        self.ps = [self.p1, self.p2, self.p3, self.p4, self.p5, self.p6, self.p7, self.p8, self.p9, self.p10,
                   self.p11, self.p12]

        self.real_name = ""

    @property
    def manufacturer_mapping(self):
        return {
            "高锰酸钾": "西安化学试剂厂",
            "重铬酸钾": "成都科龙化工试剂厂",
            "氢氧化钾": "国药集团化学试剂有限公司",
            "氧气": "",
            "氩": "陕西鑫鼎石化物资有限公司",
            "硝酸": "成都科龙化工试剂厂",
            "磷酸": "成都科龙化工试剂厂",
            "硝酸银": "西安化学试剂厂",
            "过硫酸铵": "天津市天力化学试剂有限公司",
            "可溶性淀粉": "西安延河化工厂",
            "硫酸亚铁": "成都科龙化工试剂厂",
            "氯化钡": "天津市致远化学试剂有限公司",
            "柠檬酸铵": "天津欧博凯化工有限公司",
            "硫酸钠": "成都科龙化工试剂厂",
            "硝酸钾": "西安化学试剂厂",
            "硝酸铵": "郑州太尼化学试剂有限公司",
            "水合氯醛": "成都科龙化工试剂厂",
            "氯化亚锡二水合物": "西安中信精细化工有限责任公司",
            "甲基橙": "天津津东天正精细化学试剂厂",
            "酚酞": "天津津东天正精细化学试剂厂",
            "酸性橙7": "国药集团化学试剂有限公司",
            "铬黑T": "北京化工厂",
            "乙二胺四乙酸": "yvo@tec"
        }

    def add_paragraph(self, content='', style=None):
        paragraph = self.new_doc.add_paragraph(content, style)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(2)
        paragraph_format.space_after = Pt(2)
        paragraph_format.line_spacing = Pt(12)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        return paragraph

    def add_heading(self, content='', level=1):
        style = "Title" if level == 0 else "Heading %d" % level
        paragraph = self.add_paragraph(content, style)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(2)
        paragraph_format.space_after = Pt(2)
        paragraph_format.line_spacing = Pt(16)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

    def parse_split_text(self, text: str):
        s_text = text.split('：')
        data_list = []
        # len(s_text) > 2
        last_key = None
        data_dict = {}
        for index, v in enumerate(s_text):
            v = v.strip()
            last_period_index = v.rfind('。')
            if last_period_index > -1 and last_period_index != len(s_text) - 1:
                _text_with_period = v.rsplit('。', maxsplit=1)
                if index == 0:
                    data_list.append(_text_with_period[0])
                    last_key = _text_with_period[1]
                    continue
                elif 0 < index < len(s_text) - 1:
                    data_dict[last_key] = _text_with_period[0]
                    last_key = _text_with_period[1]
                else:   # last one
                    data_dict[last_key] = _text_with_period[1]
                    break
            else:
                if index != len(s_text) - 1:
                    raise Exception("Cannot find key by split `。`")
                data_dict[last_key] = v
                break
        return data_list, data_dict

    def parse_runs(self, runs: List[Run], add_heading_level=7):
        m = {}
        writing = False
        title = ''
        data: List[str] = []
        for run in runs:
            if run.text == ' ' or not run.text:
                continue

            if '：' in run.text:
                _text = run.text.split('：')
                if len(_text) == 2:
                    if not _text[1]:
                        title = run.text.split('：')[0]
                        writing = True
                        continue
                    else:
                        if _text[0].rfind('。') > -1:
                            _text_with_period = _text[0].rsplit('。', maxsplit=1)
                            data.append(_text_with_period[0]+"。")
                            title = _text_with_period[1]
                            writing = True
                        else:
                            title = _text[0]
                        m[title] = _text[1]
                        continue
                elif len(_text) > 2:
                    _l, _d = self.parse_split_text(run.text.strip())
                    data += _l
                    m.update(_d)
                    continue

            if writing and title:
                m[title] = run.text.strip().strip('  ')
            else:
                data.append(run.text.strip())

        if m:
            if data:
                self.add_paragraph("".join(data))
            for title, content in m.items():
                self.add_heading(title, level=add_heading_level)
                self.add_paragraph(content)
        else:
            # 没有`：`标题分割，意味着只有一块文字
            content = " ".join(data)
            self.add_paragraph(content)

    def _p_name_para(self, p: Paragraph, yvo_p: P):
        self.add_heading(f'第{yvo_p.num}部分 {yvo_p.name}', 2)
        full_name = " ".join([run.text for run in p.runs if run.text and run.text != ' '])
        for _p in yvo_p.sub_p:
            _t = full_name.split("：")
            next_is_yes = False
            for ii in _t:
                if next_is_yes:
                    self.add_heading(_p.name, level=6)
                    content_split = ii.rsplit(' ', maxsplit=1)
                    if len(content_split) == 2:
                        _para_new = self.add_paragraph(content_split[0].strip())
                        if _p.name == '化学品中文名':
                            self.real_name = _para_new.text
                    elif len(content_split) == 1:
                        for _pp in yvo_p.sub_p:
                            if _pp.name in content_split[0]:
                                content_split[0] = content_split[0].replace(_pp.name, '')
                        _para_new = self.add_paragraph(content_split[0].strip())
                        if _p.name == '化学品中文名':
                            self.real_name = _para_new.text
                    _p.done()
                    break

                if _p.name in ii:
                    next_is_yes = True
        if yvo_p.is_done:
            self.add_heading("生产厂商", level=6)
            self.add_paragraph(self.manufacturer_mapping.get(self.real_name))

    def _p_title_multi_para(self, p: Paragraph, yvo_p: P):
        """
        标题one para，还有列表小标题
        :param p:
        :param yvo_p:
        :return:
        """
        self._p_one_para(p, yvo_p)
        yvo_p.process_type = ProcessType.multi_para
        yvo_p.is_done = False
        yvo_p.waiting = True
        return LoopControl.continue_

    def _p_one_para(self, p: Paragraph,
                    yvo_p: P):
        self.add_heading(f'第{yvo_p.num}部分 {yvo_p.name}', 2)
        for i, r in enumerate(p.runs):
            if yvo_p.one_para_contain_p_name and yvo_p.name in r.text:
                if r.text.strip()[len(yvo_p.old_startswith):] == '':
                    self.parse_runs(p.runs[i+1:])
                else:
                    self.parse_runs(p.runs)
                yvo_p.done()
                break
            else:
                self.parse_runs(p.runs)
                yvo_p.done()
                break

    def _p_multi_para(self, p: Paragraph, yvo_p) -> LoopControl:
        for sub_p in yvo_p.sub_p:
            if p.text.strip().startswith(sub_p.name):
                sub_p.waiting = True
                if self.multi_para_multi_runs_p and not self.multi_para_multi_runs_p.is_done:
                    self.multi_para_multi_runs_p.done()
                self.multi_para_multi_runs_p = sub_p
                self.add_heading(sub_p.name, 3)
                return LoopControl.continue_

            if sub_p.waiting:
                self.parse_runs(runs=p.runs)
                if not sub_p.multi_para_multi_runs:
                    sub_p.done()

    def find_first_un_done(self) -> Optional[P]:
        for i in self.ps:
            if not i.is_done:
                return i

    def done_last_multi_runs_p(self) -> P:
        if self.multi_para_multi_runs_p and not self.multi_para_multi_runs_p.is_done:
            self.multi_para_multi_runs_p.done()
        _p = self.find_first_un_done()
        if not _p:
            raise Exception('Done')
        return _p

    def parse(self):
        for p in iter_block_items(self.old_doc):
            my_process = self.find_first_un_done()
            if not my_process:
                print(f'{self.real_name} finished!')
                break
            if isinstance(p, Paragraph):
                if p.style.name.startswith('Heading'):
                    my_process = self.done_last_multi_runs_p()
                if p.text.startswith(my_process.old_startswith):    # important
                    my_process = self.done_last_multi_runs_p()
                    if my_process.process_type is ProcessType.name_para:
                        self._p_name_para(p, my_process)
                        continue

                    if my_process.process_type is ProcessType.one_para:
                        if my_process.need_wait:
                            my_process.waiting = True
                            continue
                        else:
                            self._p_one_para(p, my_process)
                            continue
                    elif my_process.process_type in {
                        ProcessType.multi_para, ProcessType.one_multi_para, ProcessType.table
                    }:
                        self.add_heading(f'第{my_process.num}部分 {my_process.name}', 2)
                        my_process.waiting = True
                        continue
                    # elif my_process.process_type is ProcessType.table:
                    #     self.new_doc.add_heading(f'第{my_process.num}部分 {my_process.name}', 1)
                    #     my_process.waiting = True
                    #     continue
                    # elif my_process.process_type is ProcessType.one_multi_para:
                    #     self.new_doc.add_heading(f'第{my_process.num}部分 {my_process.name}', 1)
                    #     my_process.waiting = True
                    #     continue
                elif my_process.waiting:
                    if my_process.process_type is ProcessType.multi_para:
                        control = self._p_multi_para(p, my_process)
                        if control is LoopControl.continue_:
                            continue
                    if my_process.process_type is ProcessType.one_para:
                        self._p_one_para(p, my_process)
                        continue
                    if my_process.process_type is ProcessType.one_multi_para:
                        self._p_title_multi_para(p, my_process)
                        continue
            elif my_process.process_type is ProcessType.table and my_process.waiting and isinstance(p, Table):
                if p.row_cells(0)[0].text.strip().startswith(my_process.table_first_cell_text):
                    _p_table_p = self.add_paragraph()
                    table = self.new_doc.add_table(rows=len(p.rows), cols=len(p.columns))
                    for x in range(len(p.rows)):
                        for y in range(len(p.columns)):
                            cell = table.cell(x, y)
                            cell.text = p.cell(x, y).text.strip()

                    for row in table.rows:
                        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                        row.height = Pt(16)
                    # self.new_doc._element._insert_tbl(p._tbl)
                    # _p_table_p.add_run().element.addnext(p._tbl)  # paragraph is not paragraph._p
                    my_process.done()
                    continue
        self.save()

    def save(self):
        self.new_doc.save(self.new_doc_path/self.old_doc_name)


if __name__ == '__main__':
    for __p in Path('docx/').glob('*.docx'):
        YvoParser(old_doc=__p.absolute(), new_doc_path=Path('sds_v2/')).parse()
