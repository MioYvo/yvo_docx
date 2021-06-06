# __author__ = "Mio"
# __email__: "liurusi.101@gmail.com"
# created: 6/4/21 10:05 PM
import shutil
from dataclasses import dataclass
from os.path import basename
from pathlib import Path
from typing import Optional, Callable, Union

from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx2txt import docx2txt

from docx import Document, ImagePart
from docx.document import Document as T_Document
from docx.text.run import Run
from docx.oxml.ns import qn

from utils.block import iter_block_items


class DoneExp(Exception):
    pass


@dataclass
class Process:
    func: Callable = None
    starts_flag: str = ""
    table_first_cell_text: str = ""
    contain_heading: set = None
    add_break_before_h1: bool = True
    add_para_break_before_h1: bool = False

    is_waiting: bool = False
    is_done: bool = False
    started: bool = False

    def done(self):
        self.is_waiting = False
        self.is_done = True
        self.started = True


class YvoParser:
    def __init__(self,
                 old_doc=Path()/'重铬酸钾__Potassium dichromate__7778-50-9.docx',
                 new_doc_path: Path = Path('sds_v1/')):
        self.old_doc: T_Document = Document(old_doc)
        self.old_doc_name = old_doc.name
        self.old_doc_path = old_doc
        self.new_doc: T_Document = Document()
        sections = self.new_doc.sections
        for section in sections:
            section.top_margin = Cm(0.5)
            section.bottom_margin = Cm(0.5)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)

        styles = ['Normal', 'Title']
        for i in range(1, 10):
            styles.append(f'Heading {i}')
        font_name = 'Source Han Sans CN'
        for style in styles:
            self.new_doc.styles[style].font.name = font_name
            self.new_doc.styles[style]._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            self.new_doc.styles[style]._element.rPr.rFonts.set(qn('w:eastAsiaTheme'), font_name)
            self.new_doc.styles[style]._element.rPr.rFonts.set(qn('w:asciiTheme'), font_name)

        self.new_doc_path = new_doc_path
        self.new_doc_path.mkdir(exist_ok=True)

        _heading = self.new_doc.add_heading('化学品安全技术说明书', 0)
        # print(_heading.style.element.xml)
        self.real_name = ""

        self.ps = [
            Process(self.p1, starts_flag="第1部分 化学品及企业标识"),
            Process(self.p2, starts_flag="第2部分 危险性概述"),
            Process(self.p_pure_chart, starts_flag="第3部分 成分/组成信息", table_first_cell_text="组分"),
            Process(self.p_combine_mini_list, starts_flag="第4部分 急救措施", add_para_break_before_h1=True),
            Process(self.p_normal_pure_text, starts_flag="第5部分 消防措施", add_break_before_h1=False),
            Process(self.p_normal_pure_text, starts_flag="第6部分 泄露应急处理"),
            Process(self.p_normal_pure_text, starts_flag="第7部分 操作处置与储存"),
            Process(self.p_normal_pure_text, starts_flag="第8部分 接触控制/个体防护"),
            Process(self.p_pure_chart, starts_flag="第9部分 理化特性"),
            Process(self.p_one_text_contain_flag, starts_flag="第10部分 稳定性和反应性", add_para_break_before_h1=True),
            Process(self.p_normal_pure_text, starts_flag="第11部分 毒理学信息", add_break_before_h1=False),
            Process(self.p_normal_pure_text, starts_flag="第12部分 生态学信息"),
            Process(self.p_normal_pure_text, starts_flag="第13部分 废弃处置"),
            Process(self.p14, starts_flag="第14部分 运输信息"),
            Process(self.p_normal_pure_text, starts_flag="第15部分 法规信息"),
            Process(self.p_normal_pure_text, starts_flag="第16部分 其他信息", contain_heading={"编写和修订信息"}),
            # Process(self.p4, starts_flag="第6部分"),
            # Process(self.p4, starts_flag="第7部分"),
            # Process(self.p4, starts_flag="第8部分"),
            # Process(self.p4, starts_flag="第9部分"),
            # Process(self.p4, starts_flag="第10部分"),
            # Process(self.p4, starts_flag="第11部分"),
            # Process(self.p4, starts_flag="第12部分"),
            # Process(self.p4, starts_flag="第13部分"),
            # Process(self.p4, starts_flag="第14部分"),
            # Process(self.p4, starts_flag="第15部分"),
            # Process(self.p4, starts_flag="第16部分"),
        ]

        self.HEADING_LEVEL_1 = 2
        self.HEADING_LEVEL_2 = 3
        self.HEADING_LEVEL_3 = 5

        self.image_file_rel = {}
        self.image_file_mapping()

        self.IMG_FILES_PATH = Path('image_tmp/')
        if self.IMG_FILES_PATH.exists():
            shutil.rmtree(self.IMG_FILES_PATH)
        self.IMG_FILES_PATH.mkdir(exist_ok=False)
        self.image_files_extract()

        self.current_process: Optional[Process] = None

    @property
    def manufacturer_mapping(self):
        return {
            "高锰酸钾": "西安化学试剂厂",
            "重铬酸钾": "成都科龙化工试剂厂",
            "氢氧化钾": "国药集团化学试剂有限公司",
            "氧气": "",
            "氩": "陕西鑫鼎石化物资有限公司",
            "硝酸": "成都科龙化工试剂厂",
            "磷酸": "成都科龙化工试剂厂",
            "硝酸银": "西安化学试剂厂",
            "过硫酸铵": "天津市天力化学试剂有限公司",
            "可溶性淀粉": "西安延河化工厂",
            "硫酸亚铁": "成都科龙化工试剂厂",
            "氯化钡": "天津市致远化学试剂有限公司",
            "柠檬酸铵": "天津欧博凯化工有限公司",
            "硫酸钠": "成都科龙化工试剂厂",
            "硝酸钾": "西安化学试剂厂",
            "硝酸铵": "郑州太尼化学试剂有限公司",
            "水合氯醛": "成都科龙化工试剂厂",
            "氯化亚锡二水合物": "西安中信精细化工有限责任公司",
            "甲基橙": "天津津东天正精细化学试剂厂",
            "酚酞": "天津津东天正精细化学试剂厂",
            "酸性橙7": "国药集团化学试剂有限公司",
            "铬黑T": "北京化工厂",
            "乙二胺四乙酸": "yvo@tec"
        }

    def add_paragraph(self, content='', style=None):
        paragraph = self.new_doc.add_paragraph(content, style)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(2)
        paragraph_format.space_after = Pt(2)
        paragraph_format.line_spacing = Pt(13)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        return paragraph

    def add_heading(self, content='', level=1, add_break_before_h1=True, filer_exists_run=True,
                    add_para_break_before_h1=False):
        if level == self.HEADING_LEVEL_1 and add_break_before_h1:
            if add_para_break_before_h1:
                self.add_paragraph()
            else:
                self.new_doc_last_run(filer_exists_run=filer_exists_run).add_break()

        style = "Title" if level == 0 else "Heading %d" % level
        paragraph = self.add_paragraph(content, style)
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(2)
        paragraph_format.space_after = Pt(2)
        paragraph_format.line_spacing = Pt(16)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        if level == self.HEADING_LEVEL_1:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def add_run(self, content='', style=None, paragraph: Paragraph = None, bold=False, underline=False, italic=False):
        if not paragraph:
            paragraph = self.new_doc_last_par(filer_exists_run=False)
            if paragraph.style.name.startswith('Heading'):
                paragraph = self.add_paragraph()
        _run = paragraph.add_run(text=content, style=style)
        _run.bold = bold
        _run.underline = underline
        _run.italic = italic

    def find_first_un_done(self) -> Optional[Process]:
        for i in self.ps:
            if not i.is_done:
                return i
        else:
            return None
            # raise DoneExp()

    def p1(self, p: Paragraph, process: Process):
        text = p.text.strip()
        self.add_heading(process.starts_flag, level=self.HEADING_LEVEL_1, add_break_before_h1=False,
                         add_para_break_before_h1=process.add_para_break_before_h1)
        zh_name_start_flag = '化学品中文名：'
        en_name_start_flag = '化学品英文名：'
        zh_name_start_index = text.find(zh_name_start_flag)
        en_name_start_index = text.find(en_name_start_flag)
        zh_name_end_index = en_name_start_index
        en_name_end_index = text.find('产品编号：')
        zh_name = text[zh_name_start_index + len(zh_name_start_flag):zh_name_end_index].strip()
        en_name = text[en_name_start_index + len(en_name_start_flag):en_name_end_index].strip()
        self.real_name = zh_name

        mapping = {
            zh_name_start_flag: zh_name,
            en_name_start_flag: en_name,
            "生产厂商：": self.manufacturer_mapping.get(self.real_name)
        }

        p = self.add_paragraph()
        for index, key in enumerate(mapping):
            p_name: Run = p.add_run(key)
            p_name.bold = True
            p_value = p.add_run(mapping[key])
            # -- line break
            if index != len(mapping) - 1:
                p_value.add_break()
            # -- line break --

        process.done()

    def image_file_mapping(self):
        for r in self.old_doc.part.rels.values():
            if isinstance(r._target, ImagePart):
                self.image_file_rel[r.rId] = basename(r._target.partname)

    def image_files_extract(self):
        docx2txt.process(self.old_doc_path, self.IMG_FILES_PATH)

    def p2(self, p: Paragraph, process: Process):
        text = p.text.strip()
        self.current_process = process
        if text.startswith(process.starts_flag):
            self.add_heading(process.starts_flag, level=self.HEADING_LEVEL_1,
                             add_break_before_h1=process.add_break_before_h1,
                             add_para_break_before_h1=process.add_para_break_before_h1)
            process.is_waiting = True
            process.is_done = False
            return
        if text.startswith('紧急情况概述'):
            self.add_heading('紧急情况概述', level=self.HEADING_LEVEL_2)
            process.is_waiting = True
            return

        if text.startswith('GHS危险性类别'):
            process.is_waiting = False
            return

        if text.startswith('标签要素'):
            self.add_heading('标签要素：', level=self.HEADING_LEVEL_2)
            new_p = self.add_paragraph()
            mini_list = text.split('      ')[1:]
            for t_text in mini_list:
                bold_k, normal_v = t_text.split('：')
                if bold_k:
                    new_r = new_p.add_run(bold_k.strip() + ': ')
                    new_r.bold = True
                    if normal_v:
                        new_r = new_p.add_run(normal_v.strip())
                        new_r.add_break()
                    # else:
                    #     new_r.add_break()

            new_p.add_run('象形图：').bold = True
            process.is_waiting = True
            # r = p.add_run()
            # r.add_picture()
            return
        if text.startswith('危险性说明'):
            process.is_waiting = False
            return
        if text.startswith('防范说明'):
            process.is_waiting = False
            self.add_heading(text, level=self.HEADING_LEVEL_2, add_break_before_h1=False)
            return
        if text.startswith('废弃处置'):
            process.is_waiting = True
            self.add_heading(text, level=self.HEADING_LEVEL_3)
            return

        if process.is_waiting:
            if 'graphicData' in p._p.xml:
                _new_p = self.new_doc.add_paragraph()
                for rId in self.image_file_rel:
                    if rId in p._p.xml:
                        image = self.IMG_FILES_PATH/self.image_file_rel[rId]
                        if 'graphicData' in self.new_doc_last_par()._p.xml:
                            last_run = self.new_doc_last_run()
                            last_run.add_picture(str(image.absolute()))
                        else:
                            _new_r = _new_p.add_run()
                            _new_r.add_picture(str(image.absolute()))
            else:
                _new_p = self.add_paragraph(text)

    def p_one_text_contain_flag(self, p: Paragraph, process: Process):
        text = p.text.strip()
        self.add_heading(process.starts_flag, level=self.HEADING_LEVEL_1, filer_exists_run=False,
                         add_break_before_h1=process.add_break_before_h1,
                         add_para_break_before_h1=process.add_para_break_before_h1)
        text = text[len(process.starts_flag):]

        new_p = self.add_paragraph()
        tiny_list = text.split('   ')
        for t_text in tiny_list:
            bold_k, normal_v = t_text.split('：')
            if bold_k:
                new_r = new_p.add_run(bold_k.strip() + ': ')
                new_r.bold = True
                if normal_v:
                    new_r = new_p.add_run(normal_v.strip())
                    new_r.add_break()
                else:
                    new_r.add_break()

        process.done()

    def p14(self, p: Paragraph, process: Process):
        text = p.text.strip()
        self.current_process = process
        if text.startswith(process.starts_flag):
            self.add_heading(process.starts_flag, level=self.HEADING_LEVEL_1, filer_exists_run=False,
                             add_break_before_h1=process.add_break_before_h1,
                             add_para_break_before_h1=process.add_para_break_before_h1)
            text = text[len(process.starts_flag):]

            new_p = self.add_paragraph()
            tiny_list = text.split('   ')
            for t_index, t_text in enumerate(tiny_list):
                bold_k, normal_v = t_text.split('：', maxsplit=1)
                if bold_k:
                    new_r = new_p.add_run(bold_k.strip() + ': ')
                    new_r.bold = True
                    if normal_v:
                        new_r = new_p.add_run(normal_v.strip())
                        if t_index != len(tiny_list) - 1:
                            new_r.add_break()
                    else:
                        if t_index != len(tiny_list) - 1:
                            new_r.add_break()
            process.is_waiting = True
            process.is_done = False
            return

        if process.is_waiting:
            if p.style.name.startswith('Heading'):
                if process.contain_heading and not self.contain_heading(text, contained_set=process.contain_heading):
                    process.is_waiting = False
                    return
                else:
                    process.is_waiting = True
                    if p.style.name == 'Heading 3':
                        self.add_heading(text, level=self.HEADING_LEVEL_2)
                    else:
                        # origin_heading = int(p.style.name.split(' ')[1])
                        # _heading = self.HEADING_LEVEL_2 + (self.HEADING_LEVEL_2 - origin_heading)
                        self.add_heading(text, level=self.HEADING_LEVEL_3)
                    return

            if p.style.name == 'Nomral':
                self.add_run(text)
            else:
                _new_p = self.add_paragraph(text)

    def p_pure_chart(self, p: Union[Paragraph, Table], process: Process):
        if isinstance(p, Paragraph):
            text = p.text.strip()
            self.current_process = process
            if text.startswith(process.starts_flag):
                self.add_heading(process.starts_flag, level=self.HEADING_LEVEL_1,
                                 add_break_before_h1=process.add_break_before_h1,
                                 add_para_break_before_h1=process.add_para_break_before_h1)
                process.is_waiting = True
                process.is_done = False
                return
        else:
            if p.row_cells(0)[0].text.strip().startswith(process.table_first_cell_text):
                self.insert_table(table=p)
                process.done()

    def p_combine_mini_list(self, p: Paragraph, process: Process):
        text = p.text.strip()
        self.current_process = process
        if text.startswith(process.starts_flag):
            self.add_heading(process.starts_flag, level=self.HEADING_LEVEL_1, filer_exists_run=False,
                             add_break_before_h1=process.add_break_before_h1,
                             add_para_break_before_h1=process.add_para_break_before_h1)
            process.is_waiting = True
            process.is_done = False
            return

        if process.is_waiting:
            new_p = self.add_paragraph()
            mini_list = text.split('      ')
            for m_list in mini_list:
                tiny_list = m_list.split('   ')
                for t_text in tiny_list:
                    bold_k, normal_v = t_text.split('：')
                    if bold_k:
                        new_r = new_p.add_run(bold_k + ': ')
                        new_r.bold = True
                        if normal_v:
                            new_r = new_p.add_run(normal_v)
                            new_r.add_break()
                        else:
                            new_r.add_break()

    def contain_heading(self, cur, contained_set):
        for head in contained_set:
            if cur.startswith(head):
                return True
        else:
            return False

    def p_normal_pure_text(self, p: Union[Paragraph, Table], process: Process):
        if isinstance(p, Paragraph):
            text = p.text.strip()
            self.current_process = process
            if text.startswith(process.starts_flag):
                self.add_heading(process.starts_flag, level=self.HEADING_LEVEL_1, filer_exists_run=False,
                                 add_break_before_h1=process.add_break_before_h1,
                                 add_para_break_before_h1=process.add_para_break_before_h1)
                process.is_waiting = True
                process.is_done = False
                return

            if p.style.name.startswith('Heading'):
                if process.contain_heading and not self.contain_heading(text, contained_set=process.contain_heading):
                    process.is_waiting = False
                    return
                else:
                    process.is_waiting = True
                    if p.style.name == 'Heading 3':
                        self.add_heading(text, level=self.HEADING_LEVEL_2)
                    else:
                        # origin_heading = int(p.style.name.split(' ')[1])
                        # _heading = self.HEADING_LEVEL_2 + (self.HEADING_LEVEL_2 - origin_heading)
                        self.add_heading(text, level=self.HEADING_LEVEL_3)
                    return

        if process.is_waiting:
            if isinstance(p, Paragraph):
                text = p.text.strip()
                if p.style.name == 'Nomral':
                    self.add_run(text)
                else:
                    _new_p = self.add_paragraph(text)
            else:
                self.insert_table(table=p)

    def insert_table(self, table: Table):
        new_table = self.new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        for x in range(len(table.rows)):
            for y in range(len(table.columns)):
                cell = new_table.cell(x, y)
                cell.text = table.cell(x, y).text.strip()

        for row in new_table.rows:
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
            row.height = Pt(16)

    def new_doc_last_par(self, filer_exists_run=True) -> Paragraph:
        if filer_exists_run:
            for p in self.new_doc.paragraphs[::-1]:
                if p.runs:
                    return p
        else:
            return self.new_doc.paragraphs[-1]

    def new_doc_last_run(self, filer_exists_run=True) -> Run:
        return self.new_doc_last_par(filer_exists_run).runs[-1]

    def done_last_process(self) -> Process:
        if self.current_process and not self.current_process.is_done:
            self.current_process.done()
            # print('-'*30)
        return self.find_first_un_done()

    def _parse(self):
        for p in iter_block_items(self.old_doc):
            process = self.find_first_un_done()
            # print(type(p), p.style.name, p.text if hasattr(p, 'text') else p)
            # if isinstance(p, Paragraph):
            #     if p.text.startswith('标签要素：   '):
            #         print(p)
            if not process:
                continue
            if isinstance(p, Paragraph):
                if p.style.name.startswith('Heading 2'):
                    process = self.done_last_process()

                if not process:
                    continue

                if p.text.strip().startswith(process.starts_flag):
                    process.started = True
                    process.func(p, process)
                    continue

                if not process.is_done and process.started:
                    process.func(p, process)
            elif isinstance(p, Table) and not process.is_done and process.started:
                process.func(p, process)
            else:
                pass

    def parse(self):
        try:
            self._parse()
        except DoneExp:
            self.save()
        except Exception as e:
            print(e)
        else:
            self.save()

        shutil.rmtree(self.IMG_FILES_PATH)

    def save(self):
        self.new_doc.save(self.new_doc_path/self.old_doc_name)


if __name__ == '__main__':
    total = list(Path('docx/').glob('*.docx'))
    for _index, __p in enumerate(total, start=1):
        YvoParser(old_doc=__p.absolute(), new_doc_path=Path('/run/media/mio/Data/sds_v3/')).parse()
        print(f"{_index}/{len(total)} {__p.name} done")
