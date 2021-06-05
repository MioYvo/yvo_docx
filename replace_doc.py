# __author__ = "Mio"
# __email__: "liurusi.101@gmail.com"
# created: 5/20/21 4:55 PM
from pathlib import Path
from docx import Document

Path('docx').mkdir(exist_ok=True)
for file in Path('output_docx').glob('*.docx'):
    doc = Document(file)
    for i in doc.paragraphs:
        # print(i.text)
        cn_index = i.text.find('化学品中文名')
        en_index = i.text.find('化学品英文名')
        if cn_index >= 0:
            cn_name = i.text[cn_index:].split('   ')[0].split('： ')[1]
        # if en_index >= 0:
        #     ene_name = i.text[en_index:].split('   ')[0].split('： ')[1]

            for t in doc.tables:
                if t.rows[0].cells[0].text == '组分':
                    en_name = t.rows[1].cells[0].text
                    cas_no = t.rows[1].cells[2].text
                    print(cn_name, en_name, cas_no)
                    name = f'{file.stem}__{en_name}__{cas_no}{file.suffix}'
                    doc.save(Path('docx') / name)

    # for i in doc.paragraphs:
    #     print(i.text)
    #     if i.text.lower().find('xixisys') >= 0:
    #         i.text = i.text.replace("XiXisys.com 免费提供，仅供参考。", " ")
    #         i.text = i.text.replace(" 如有疑问，请联系 sds@xixisys.com 咨询。", " ")

    # for i in doc.paragraphs:
    #     # print(i.text)
    #     cn_index = i.text.find('化学品中文名')
    #     en_index = i.text.find('化学品英文名')
    #     if cn_index >= 0:
    #         cn_name = i.text[cn_index:].split('   ')[0].split('： ')[1]
    #     if en_index >= 0:
    #         ene_name = i.text[en_index:].split('   ')[0].split('： ')[1]
    #     break

    # doc.save(f"{Path('output_docx')/file.stem}.docx")


# doc = Document('乙二胺四乙酸.docx')
# for i in doc.paragraphs:
#     cn_index = i.text.find('化学品中文名')
#     en_index = i.text.find('化学品英文名')
#     if cn_index >= 0:
#         print(i.text[cn_index:].split('   ')[0].split('： ')[1])
#     if en_index >= 0:
#         print(i.text[en_index:].split('   ')[0].split('： ')[1])
