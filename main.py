# __author__ = "Mio"
# __email__: "liurusi.101@gmail.com"
# created: 5/19/21 11:34 PM
from pathlib import Path

from docx import Document
from htmldocx import HtmlToDocx
import httpx
from lxml import etree


def print_dir_file():
    print('-' * 30)
    for i in Path().glob('output_docx/*'):
        print(i)
    print('-' * 30)


def get_data(url: str):
    document = Document()
    new_parser = HtmlToDocx()
    res = httpx.get(url)
    content = res.content.decode()
    content = content.replace("XiXisys.com 免费提供，仅供参考。", " ")
    content = content.replace(" 如有疑问，请联系 sds@xixisys.com 咨询。", " ")
    new_parser.add_html_to_document(content, document)
    name = etree.HTML(res.content).xpath('/html/body/article/section[1]/div[1]/span[2]')[0].text
    # document.save(f'output_docx/{name}.docx')

    for i in document.paragraphs:
        # print(i.text)
        cn_index = i.text.find('化学品中文名')
        en_index = i.text.find('化学品英文名')
        if cn_index >= 0:
            cn_name = i.text[cn_index:].split('   ')[0].split('： ')[1]
        # if en_index >= 0:
        #     ene_name = i.text[en_index:].split('   ')[0].split('： ')[1]

            for t in document.tables:
                if t.rows[0].cells[0].text == '组分':
                    en_name = t.rows[1].cells[0].text
                    cas_no = t.rows[1].cells[2].text
                    print(cn_name, en_name, cas_no)
                    name = f'{cn_name}__{en_name}__{cas_no}.docx'
                    document.save(Path('docx') / name)
    return name


if __name__ == '__main__':
    new_added = None
    Path('output_docx').mkdir(exist_ok=True)
    while True:
        print_dir_file()
        if new_added:
            print(f'{new_added} added')
        try:
            url = input('输入 xixisys-api 地址: ')
        except KeyboardInterrupt:
            print('Bye baby...')
            break
        except Exception:
            break
        new_added = get_data(url)
